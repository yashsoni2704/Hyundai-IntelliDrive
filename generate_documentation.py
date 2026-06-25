"""
Generate comprehensive Word documentation for Hyundai Knowledge Assistant.

Run from project root:
    pip install python-docx
    python generate_documentation.py

Output: Hyundai_Knowledge_Assistant_Project_Documentation.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT = "Hyundai_Knowledge_Assistant_Project_Documentation.docx"
VERSION = "2.2.0"
FAQ_COUNT = 165


def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HYUNDAI KNOWLEDGE ASSISTANT\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 44, 95)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Complete Technical Documentation\n")
    r.font.size = Pt(16)
    r.bold = True

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub2.add_run(
        "Retrieval-Based Semantic FAQ Chatbot · Auth · OTP Email · Bookings · Admin Panel\n\n"
    )
    r2.font.size = Pt(12)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Prepared for: Project Presentation, Viva & Code Review\n").font.size = Pt(11)
    info.add_run(f"Version: {VERSION}\n").font.size = Pt(11)
    info.add_run(
        "Stack: React 18 + Vite 6 + FastAPI + ChromaDB + SQLite + BGE-M3 Embeddings"
    ).font.size = Pt(11)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "002C5F")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)

    add_heading(doc, "Table of Contents", 1)
    for item in [
        "1. Project Overview",
        "2. Complete End-to-End Flow",
        "3. System Architecture",
        "4. Technology Stack & Libraries",
        "5. Knowledge Base Design & Chunking Strategy",
        "6. Embeddings & Semantic Search (ChromaDB)",
        "7. Context, Sessions & Clarification Logic",
        "8. Database Design (SQLite)",
        "9. API Endpoints (Complete)",
        "10. Authentication, OTP & Email",
        "11. Test Drive Booking System",
        "12. Admin Dashboard",
        "13. Frontend Architecture",
        "14. Every File Explained",
        "15. Code Comments Guide (How to Read the Source)",
        "16. Environment Variables",
        "17. How to Run",
        "18. Viva Questions & Answers",
        "19. Limitations & Future Scope",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    # 1. OVERVIEW
    add_heading(doc, "1. Project Overview", 1)
    add_para(
        doc,
        "Hyundai Knowledge Assistant is a full-stack web application for Hyundai showroom customers. "
        "It answers vehicle questions from a verified FAQ knowledge base, books test drives, and "
        "supports user accounts with email OTP security.",
    )
    add_para(doc, "Critical design choice: This is a RETRIEVAL-BASED chatbot.", bold=True)
    add_para(
        doc,
        "It does NOT use ChatGPT, Gemini, or any LLM to generate text. Every answer is copied "
        "verbatim from the Excel FAQ database after semantic matching. If no match is found, "
        "the bot says 'Sorry, no data found.' — it never invents prices or specs.",
    )
    add_para(doc, "Key features:", bold=True)
    for f in [
        f"Semantic FAQ search — {FAQ_COUNT} Q&A pairs, ChromaDB + BGE-M3 embeddings",
        "Topic + vehicle re-ranking (price, mileage, seats, compare)",
        "Session context — follow-ups like 'its mileage' after asking about Creta",
        "Clarification prompts — asks car model when user says 'its price' with no context",
        "Spelling tolerance — typos like milage, tucsan, compair are normalized",
        "Registration & login with email OTP (2FA)",
        "Forgot / reset password via OTP",
        "Test drive booking with 1-hour slots (10 AM–8 PM)",
        "In-chat slot table when user asks for available timings",
        "Dynamic follow-up suggestion chips",
        "Per-user chat sessions stored in SQLite",
        "Sidebar: last 5 Q&A history across all logins",
        "Admin dashboard: bookings CRUD + paginated chat monitor (IST timestamps)",
    ]:
        add_bullet(doc, f)

    # 2. END-TO-END FLOW
    add_heading(doc, "2. Complete End-to-End Flow", 1)

    add_heading(doc, "2.1 Application Startup", 2)
    for s in [
        "Developer runs: uvicorn app:app --port 8000 (backend) and npm run dev (frontend).",
        "FastAPI lifespan hook calls init_db() → creates SQLite tables if missing.",
        "vector_store.initialize() reads data/hyundai_faq.xlsx via Pandas.",
        "If Excel hash/count/model changed → embed all FAQ questions → store in ChromaDB.",
        "Frontend loads at http://127.0.0.1:5173; /api/* proxied to port 8000.",
    ]:
        add_numbered(doc, s)

    add_heading(doc, "2.2 User Login Flow", 2)
    for s in [
        "User opens AuthModal → enters email + password.",
        "POST /auth/login verifies bcrypt hash → generates 6-digit OTP → saves to otp_codes table.",
        "email_service.py sends HTML OTP email via Gmail SMTP (never returned to UI).",
        "User enters OTP → POST /auth/verify-login-otp → JWT issued.",
        "AuthContext stores token in localStorage; api.js sends Authorization: Bearer header.",
        "useSessionChat calls POST /chat/session/start → new ChatSession with empty context.",
        "Middle chat area is blank; sidebar loads GET /chat/recent (last 5 user Q&As).",
    ]:
        add_numbered(doc, s)

    add_heading(doc, "2.3 FAQ Chat Flow (Logged-In User)", 2)
    for s in [
        "User types message → ChatPage.handleSend → POST /chat with message + session_id.",
        "app.py normalizes typos (normalize_message).",
        "Loads session context: last_vehicle, last_topic, pending_topic from chat_sessions.context_json.",
        "If vague ('its price') and no car context → clarification response, pending_topic saved.",
        "If clear → resolve_query expands 'its mileage' → 'What is the mileage of Hyundai Creta?'",
        "slot_intent.py checked on ORIGINAL message (not expanded) for booking slot queries.",
        "chroma_db.search: embed query → top-12 cosine search → filter by topic + vehicle → best match.",
        "If similarity ≥ 0.55 → return stored answer from metadata.",
        "suggestions.py builds up to 4 follow-up chips.",
        "chat_log_service logs query/answer; apply_exchange_to_session updates context.",
        "Frontend displays answer + suggestions; refreshRecent updates sidebar.",
    ]:
        add_numbered(doc, s)

    add_heading(doc, "2.4 Test Drive Booking Flow", 2)
    for s in [
        "User clicks suggestion or asks 'available timings' → slot intent detected.",
        "booking_service returns next 5 free slots from SQLite bookings table.",
        "SlotTable.jsx renders slots with Book buttons.",
        "Authenticated user books → POST /bookings → unique (date, time_slot) constraint prevents double booking.",
    ]:
        add_numbered(doc, s)

    add_heading(doc, "2.5 Admin Flow", 2)
    for s in [
        "Admin logs in via /auth/admin-login → App.jsx renders AdminDashboard instead of ChatPage.",
        "Admin manages bookings, views slot grid, monitors chat logs with pagination (10 per page, newest first).",
        "Email search filter on chat logs; timestamps displayed in IST (Asia/Kolkata).",
    ]:
        add_numbered(doc, s)

    # 3. ARCHITECTURE
    add_heading(doc, "3. System Architecture", 1)
    add_para(doc, "Three-tier architecture:", bold=True)
    add_numbered(doc, "Presentation — React 18 + Vite (port 5173)")
    add_numbered(doc, "Application — FastAPI + Python (port 8000)")
    add_numbered(doc, "Data — SQLite (users, bookings, sessions, logs) + ChromaDB (vectors) + Excel (source)")
    add_para(doc, "")
    add_para(
        doc,
        "Request path: Browser → Vite proxy (/api → :8000) → FastAPI router → "
        "Service layer (chroma_db, context_service, booking_service) → SQLite/ChromaDB → JSON response.",
    )

    # 4. TECH STACK
    add_heading(doc, "4. Technology Stack & Libraries", 1)
    add_table(
        doc,
        ["Library", "Version / Role"],
        [
            ("FastAPI", "REST API framework — routes, validation, dependency injection"),
            ("Uvicorn", "ASGI server running the FastAPI app"),
            ("ChromaDB", "Vector database — persistent cosine similarity search (HNSW index)"),
            ("sentence-transformers", "Loads BAAI/bge-m3 embedding model"),
            ("BAAI/bge-m3", "1024-dim multilingual embeddings (downloaded from Hugging Face)"),
            ("Pandas + openpyxl", "Read hyundai_faq.xlsx Question/Answer columns"),
            ("SQLAlchemy 2", "ORM for SQLite — User, Booking, ChatSession, ChatLog, OtpRecord"),
            ("bcrypt", "One-way password hashing"),
            ("PyJWT", "Signed JWT tokens for authentication"),
            ("python-dotenv", "Load secrets from backend/.env"),
            ("Pydantic v2", "Request/response schema validation"),
            ("smtplib + email.mime", "Gmail SMTP for OTP emails"),
            ("React 18", "UI components — chat, modals, admin"),
            ("Vite 6", "Dev server, HMR, /api proxy"),
            ("Plain CSS", "Dark ChatGPT-style theme in index.css"),
        ],
    )

    # 5. CHUNKING
    add_heading(doc, "5. Knowledge Base Design & Chunking Strategy", 1)
    add_para(doc, "Important: This project does NOT use document chunking (no LangChain RecursiveCharacterTextSplitter, no sliding windows over PDFs).", bold=True)
    add_para(
        doc,
        "Instead we use ATOMIC FAQ-LEVEL INDEXING — each row in the Excel file is one complete "
        "knowledge unit (one Question + one Answer). This is the correct approach for structured FAQ data.",
    )
    add_para(doc, "Indexing unit (our 'chunk'):", bold=True)
    add_table(
        doc,
        ["Field", "What is stored", "Used for"],
        [
            ("document (ChromaDB)", "FAQ Question text only", "Embedding input + display document"),
            ("metadata.question", "Same question string", "Retrieval debugging"),
            ("metadata.answer", "Full answer text", "Returned to user (never embedded)"),
            ("id", "faq_0, faq_1, …", "Unique ChromaDB identifier"),
        ],
    )
    add_para(doc, "Why embed questions only (not answers)?", bold=True)
    add_para(
        doc,
        "User queries are phrased as questions. Embedding questions and matching query→question "
        "gives better semantic alignment than embedding answers (which are declarative statements).",
    )
    add_para(doc, "Ingestion pipeline (data_loader.py + chroma_db.py):", bold=True)
    for s in [
        f"Read {FAQ_COUNT} rows from Excel (columns: Question, Answer).",
        "Skip empty or NaN rows.",
        "Compute SHA-256 hash of Excel file (compute_excel_hash).",
        "Compare hash + count + embedding model vs ingestion_meta.json.",
        "If changed: delete ChromaDB collection, embed all questions via embed_texts(), add to collection.",
        "Store ingestion metadata (hash, count, model, timestamp).",
    ]:
        add_numbered(doc, s)
    add_para(doc, "Re-ingestion triggers:", bold=True)
    add_bullet(doc, "Excel file content changed (hash mismatch)")
    add_bullet(doc, "FAQ count changed")
    add_bullet(doc, "ChromaDB empty or count mismatch")
    add_bullet(doc, "EMBEDDING_MODEL env variable changed")

    # 6. EMBEDDINGS & SEARCH
    add_heading(doc, "6. Embeddings & Semantic Search (ChromaDB)", 1)
    add_para(doc, "Embedding generation (embeddings.py):", bold=True)
    for s in [
        "SentenceTransformer loads BAAI/bge-m3 once (@lru_cache).",
        "model.encode(texts, normalize_embeddings=True) → unit vectors.",
        "Cosine similarity = dot product of normalized vectors.",
    ]:
        add_bullet(doc, s)
    add_para(doc, "Search pipeline (chroma_db.py search method):", bold=True)
    for s in [
        "normalize_message(query) — fix typos.",
        "embed_query(query) → 1024-dimensional vector.",
        "ChromaDB query: n_results=12, space=cosine (distance = 1 - similarity).",
        "_pick_best_candidate: filter candidates where FAQ topic/vehicle matches query intent.",
        "Combined score: 65% similarity + 35% vehicle match.",
        "If best similarity < SIMILARITY_THRESHOLD (0.55) → 'Sorry, no data found.'",
        "Fallback: _lexical_search (keyword overlap + SequenceMatcher) if embedding model fails.",
    ]:
        add_numbered(doc, s)
    add_para(doc, "TOPIC_SIGNALS used for filtering:", bold=True)
    add_bullet(doc, "price → price, cost, lakh, rupee, starting")
    add_bullet(doc, "mileage → mileage, kmpl, fuel efficiency")
    add_bullet(doc, "seats → seat, seater, seating, capacity")
    add_bullet(doc, "compare → compare, versus, vs")

    # 7. CONTEXT
    add_heading(doc, "7. Context, Sessions & Clarification Logic", 1)
    add_para(doc, "Session context (stored as JSON in chat_sessions.context_json):", bold=True)
    add_table(
        doc,
        ["Key", "Purpose", "Example"],
        [
            ("last_vehicle", "Car user is discussing", "Creta"),
            ("last_topic", "Last topic asked", "price"),
            ("pending_topic", "Waiting for user to name car", "price (after 'its price')"),
            ("pending_vehicle", "Waiting for user to ask topic", "Tucson (after user said 'tucson')"),
            ("recent_queries", "Last 8 raw queries", "['price of creta', 'its mileage']"),
        ],
    )
    add_para(doc, "context_service.py responsibilities:", bold=True)
    for s in [
        "normalize_message — spelling fixes, vehicle alias mapping",
        "detect_vehicle / detect_topic — parse user intent",
        "is_vague_query — detects it/its/this car phrases",
        "resolve_query — expand vague follow-ups using session context",
        "needs_clarification — ask user when car model unknown",
        "prepare_clarification_context — save pending_topic for two-step dialog",
        "update_context — update last_vehicle/topic after each exchange",
    ]:
        add_bullet(doc, s)
    add_para(doc, "Example context dialog:", bold=True)
    add_numbered(doc, "User: 'What is the price of Hyundai Creta?' → context: last_vehicle=Creta, last_topic=price")
    add_numbered(doc, "User: 'its mileage' → resolved: 'What is the mileage of Hyundai Creta?'")
    add_numbered(doc, "User: 'price of alcazar' → context switches to Alcazar")
    add_numbered(doc, "User: 'compare it with creta' → 'Compare Hyundai Alcazar and Creta'")
    add_numbered(doc, "Fresh login + 'its price' → clarification: 'Which car model?' → user: 'Creta' → price answer")

    # 8. DATABASE
    add_heading(doc, "8. Database Design (SQLite)", 1)
    add_para(doc, "File: backend/app.db", bold=True)
    add_table(
        doc,
        ["Table", "Purpose", "Key columns"],
        [
            ("users", "Accounts", "id, email, password_hash, full_name"),
            ("bookings", "Test drives", "user_id, booking_date, time_slot, vehicle_model, status"),
            ("otp_codes", "Email OTPs", "email, code, purpose, expires_at, used"),
            ("chat_sessions", "Per-login chat context", "user_id, context_json, is_active"),
            ("chat_logs", "All Q&A for admin + sidebar", "query, answer, found, response_type, session_id"),
        ],
    )
    add_bullet(doc, "Unique constraint on bookings(booking_date, time_slot) prevents double booking")
    add_bullet(doc, "Chat logs store IST-formatted timestamps for admin display")

    # 9. API
    add_heading(doc, "9. API Endpoints (Complete)", 1)
    add_table(
        doc,
        ["Method", "Endpoint", "Description"],
        [
            ("GET", "/health", "Health check"),
            ("GET", "/stats", "FAQ + ChromaDB statistics"),
            ("POST", "/chat", "Main chat endpoint"),
            ("POST", "/chat/session/start", "New session on login"),
            ("POST", "/chat/session/end", "End session on logout"),
            ("GET", "/chat/session/{id}/messages", "Restore session messages"),
            ("GET", "/chat/recent", "Last 5 Q&A for sidebar"),
            ("POST", "/auth/register", "Register + send OTP"),
            ("POST", "/auth/verify-register-otp", "Complete registration"),
            ("POST", "/auth/login", "Login + send OTP"),
            ("POST", "/auth/verify-login-otp", "Complete login → JWT"),
            ("POST", "/auth/admin-login", "Admin JWT (no OTP)"),
            ("POST", "/auth/forgot-password", "Send reset OTP"),
            ("POST", "/auth/reset-password", "Reset with OTP"),
            ("GET", "/auth/me", "Current user profile"),
            ("GET", "/bookings/dates", "Bookable dates"),
            ("GET", "/bookings/slots", "Slots for a date"),
            ("POST", "/bookings", "Book test drive"),
            ("GET", "/bookings/my", "User's bookings"),
            ("GET", "/admin/bookings", "All bookings (admin)"),
            ("GET", "/admin/chat-logs?page=&per_page=&email=", "Paginated chat monitor"),
        ],
    )

    # 10-12 condensed but complete
    add_heading(doc, "10. Authentication, OTP & Email", 1)
    add_para(doc, "Passwords hashed with bcrypt. JWT contains sub, email, role, exp. OTP: 6 digits, 10 min expiry, single use, sent via Gmail SMTP only.")

    add_heading(doc, "11. Test Drive Booking System", 1)
    add_para(doc, "Slots: 10:00–20:00 hourly, 14 days ahead. slot_intent.py detects availability queries. Conflict returns next free slot suggestion.")

    add_heading(doc, "12. Admin Dashboard", 1)
    add_para(doc, "AdminDashboard.jsx: stats cards, booking CRUD, slot grid, chat monitor with pagination (page 1 = newest), email search, IST times.")

    # 13. FRONTEND
    add_heading(doc, "13. Frontend Architecture", 1)
    add_para(doc, "Entry: index.html → main.jsx → App.jsx", bold=True)
    add_para(doc, "App.jsx: if user.is_admin → AdminDashboard, else ChatPage.", bold=False)
    add_para(doc, "ChatPage.jsx orchestrates:", bold=True)
    for c in [
        "useSessionChat — session lifecycle, messages, sidebar recent Q&A",
        "useSuggestionTracker — tracks clicked suggestion chip IDs",
        "Sidebar — recent 5 questions, navigation",
        "ChatArea + MessageBubble — message display",
        "PastExchangeView — view historical Q&A from sidebar",
        "AuthModal — login/register/OTP/reset",
        "BookingModal — date/slot picker",
        "api.js — all HTTP calls with JWT header",
    ]:
        add_bullet(doc, c)

    # 14. FILES
    add_heading(doc, "14. Every File Explained", 1)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ("backend/app.py", "FastAPI entry — /chat pipeline, lifespan, routers"),
            ("backend/chroma_db.py", "ChromaDB vector store, search, re-ranking"),
            ("backend/context_service.py", "Session context, clarification, query resolution"),
            ("backend/embeddings.py", "BGE-M3 model load + embed_texts/embed_query"),
            ("backend/data_loader.py", "Excel read, hash, ingestion metadata"),
            ("backend/session_service.py", "ChatSession CRUD, context persistence"),
            ("backend/chat_log_service.py", "Log interactions, pagination, recent exchanges"),
            ("backend/suggestions.py", "Follow-up suggestion chips"),
            ("backend/slot_intent.py", "Detect slot/timing queries"),
            ("backend/booking_service.py", "Slot availability and booking logic"),
            ("backend/models.py", "SQLAlchemy table definitions"),
            ("backend/config.py", "Environment variables and paths"),
            ("data/hyundai_faq.xlsx", f"Source FAQ data ({FAQ_COUNT} pairs)"),
            ("frontend/src/pages/ChatPage.jsx", "Main chat UI controller"),
            ("frontend/src/hooks/useSessionChat.js", "Session + sidebar state hook"),
            ("frontend/src/services/api.js", "Backend HTTP client"),
            ("generate_documentation.py", "This documentation generator script"),
        ],
    )

    # 15. CODE COMMENTS GUIDE
    add_heading(doc, "15. Code Comments Guide (How to Read the Source)", 1)
    add_para(
        doc,
        "All core backend and frontend files now include module-level docstrings and inline comments "
        "explaining what each section does. Read files in this order to understand the project:",
    )
    for s in [
        "config.py — paths and settings",
        "models.py — database tables",
        "data_loader.py → embeddings.py → chroma_db.py — knowledge base pipeline",
        "context_service.py — query understanding",
        "app.py — ties everything together in /chat",
        "session_service.py + chat_log_service.py — persistence",
        "api.js → useSessionChat.js → ChatPage.jsx — frontend flow",
    ]:
        add_numbered(doc, s)
    add_para(doc, "Python syntax quick reference used in this project:", bold=True)
    add_table(
        doc,
        ["Syntax", "Meaning"],
        [
            ("def fn(x: str) -> dict:", "Function with type hints (x is str, returns dict)"),
            ("@app.post('/chat')", "FastAPI route decorator — handles POST requests"),
            ("Depends(get_db)", "FastAPI dependency injection — provides DB session"),
            ("Mapped[str]", "SQLAlchemy 2 typed column declaration"),
            ("list[str] | None", "Union type — optional list of strings (Python 3.10+)"),
            ("@lru_cache(maxsize=1)", "Cache function result — model loaded once"),
            ("async def / await", "Asynchronous function (used in FastAPI lifespan)"),
            ("useCallback / useEffect", "React hooks — memoized functions and side effects"),
        ],
    )

    # 16. ENV
    add_heading(doc, "16. Environment Variables", 1)
    add_table(
        doc,
        ["Variable", "Description", "Default"],
        [
            ("JWT_SECRET", "JWT signing key", "change in production"),
            ("SMTP_USER / SMTP_PASSWORD", "Gmail credentials", "required for OTP"),
            ("EMBEDDING_MODEL", "HuggingFace model", "BAAI/bge-m3"),
            ("SIMILARITY_THRESHOLD", "Min match score", "0.55"),
            ("COLLECTION_NAME", "ChromaDB collection", "hyundai_faq"),
            ("CORS_ORIGINS", "Allowed frontend URLs", "localhost:5173"),
            ("DEBUG_MODE", "Dev flags", "true"),
        ],
    )

    # 17. RUN
    add_heading(doc, "17. How to Run", 1)
    add_numbered(doc, "cd backend && pip install -r requirements.txt")
    add_numbered(doc, "Copy .env.example to .env, set SMTP + JWT")
    add_numbered(doc, "uvicorn app:app --reload --port 8000")
    add_numbered(doc, "cd frontend && npm install && npm run dev")
    add_numbered(doc, "Open http://127.0.0.1:5173")
    add_numbered(doc, "Regenerate docs: python generate_documentation.py")

    # 18. VIVA
    add_heading(doc, "18. Viva Questions & Answers", 1)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ("What chunking method is used?", "Atomic FAQ chunking — each Excel Q&A row = 1 document. No text splitting."),
            ("What is embedded — question or answer?", "Question only. Answer stored in ChromaDB metadata."),
            ("Why retrieval instead of LLM?", "No hallucination — showroom needs accurate prices/specs."),
            ("What is BGE-M3?", "Multilingual embedding model producing 1024-dim semantic vectors."),
            ("What is cosine similarity?", "Measures angle between vectors; 1 = identical meaning."),
            ("How does context work?", "JSON in chat_sessions — tracks last car, topic, pending clarifications."),
            ("What is ChromaDB HNSW?", "Approximate nearest neighbor index for fast vector search."),
            ("How are typos handled?", "normalize_message in context_service.py with SPELLING_FIXES dict."),
            ("Where is chat history?", "SQLite chat_logs (admin + sidebar); session messages for current chat."),
            ("What is JWT?", "Signed token proving identity; sent in Authorization header."),
        ],
    )

    # 19. LIMITATIONS
    add_heading(doc, "19. Limitations & Future Scope", 1)
    add_bullet(doc, "Answers limited to FAQ Excel content")
    add_bullet(doc, "English FAQ only")
    add_bullet(doc, "BGE-M3 ~2GB first-time download")
    add_bullet(doc, "SQLite for dev; PostgreSQL recommended for production")
    add_bullet(doc, "Future: Hindi FAQs, cloud deploy, rate limiting, live price API")

    doc.add_page_break()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run(f"— End of Document —\nHyundai Knowledge Assistant v{VERSION}")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 44, 95)

    return doc


if __name__ == "__main__":
    document = build_document()
    out = Path(OUTPUT)
    document.save(out)
    print(f"Documentation saved to: {out.resolve()}")
